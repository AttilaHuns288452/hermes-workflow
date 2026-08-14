# -*- coding: utf-8 -*-
"""Dump .docx structure: paragraphs (index/style/preview) + all tables.
Usage: PYTHONPATH= uv run --with python-docx python dump_docx.py <path.docx>
"""
import sys
from docx import Document

doc = Document(sys.argv[1])

print("=== BODY PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        style = p.style.name if p.style is not None else "?"
        print(f"[P{i}] ({style}) {t[:100]}")

print("\n=== TABLES ===")
for ti, tbl in enumerate(doc.tables):
    print(f"--- Table {ti}: {len(tbl.rows)} rows x {len(tbl.columns)} cols ---")
    for ri, row in enumerate(tbl.rows):
        cells = [c.text.strip().replace("\n", " / ")[:60] for c in row.cells]
        print(f"  R{ri}: {cells}")
