# -*- coding: utf-8 -*-
"""Proven helpers for in-place .docx editing with python-docx.
Copy into your edit script; see SKILL.md for the full workflow.
Run with: PYTHONPATH= uv run --with python-docx python your_script.py
"""
from copy import deepcopy
from docx.oxml.ns import qn


def set_para_text(p, text):
    """Replace a paragraph's content, preserving paragraph style and the
    first run's character formatting; '\\n' becomes real line breaks (w:br)."""
    rpr = None
    if p.runs:
        rpr = p.runs[0]._element.find(qn('w:rPr'))
        if rpr is not None:
            rpr = deepcopy(rpr)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            br = p.add_run()
            br.add_break()
        run = p.add_run(line)
        if rpr is not None:
            run._element.insert(0, deepcopy(rpr))


def delete_para(p):
    p._element.getparent().remove(p._element)


def set_cell(cell, text):
    for r in list(cell.paragraphs[0].runs):
        r._element.getparent().remove(r._element)
    cell.paragraphs[0].add_run(text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def insert_row_before(table, idx, texts):
    """Insert a new row before table row idx, copying that row's formatting."""
    new_tr = deepcopy(table.rows[idx]._tr)
    table.rows[idx]._tr.addprevious(new_tr)
    for c, txt in zip(table.rows[idx].cells, texts):
        set_cell(c, txt)


def replace_wt(doc, old, new):
    """XML-level text replace covering TOC field results and textboxes that
    doc.paragraphs cannot see. Returns number of w:t nodes changed."""
    n = 0
    for t in doc.element.iter(qn('w:t')):
        if t.text and old in t.text:
            t.text = t.text.replace(old, new)
            n += 1
    return n


def frag_replace(para, old, new):
    """Replace a fragment inside one run; returns False if the fragment is
    split across runs (rewrite the whole paragraph instead)."""
    for r in para.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return True
    return False
